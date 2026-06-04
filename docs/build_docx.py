"""Build E-Barber-System-CzescII.docx from the markdown source.

Uses python-docx to render the project's second-part documentation as a
polished Word file with title page, TOC, numbered headings, tables and
monospace code blocks. Polish characters survive through UTF-8.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

DOC_DIR = Path(__file__).resolve().parent
SRC = DOC_DIR / "E-Barber-System-CzescII.md"
DEST = DOC_DIR / "E-Barber-System-CzescII.docx"


def add_field(paragraph, instr_text: str, placeholder: str = "Pole zostanie obliczone w Wordzie (F9)."):
    """Insert a calculated Word field (e.g. TOC, PAGE)."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder_el = OxmlElement("w:t")
    placeholder_el.text = placeholder
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(placeholder_el)
    run._r.append(fld_char_end)


def ensure_styles(doc: Document) -> None:
    styles = doc.styles

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.base_style = styles["Normal"]
        code_font = code.font
        code_font.name = "Consolas"
        code_font.size = Pt(9)
        code.paragraph_format.left_indent = Cm(0.4)
        code.paragraph_format.right_indent = Cm(0.4)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.line_spacing = 1.15

    base_font = styles["Normal"].font
    base_font.name = "Calibri"
    base_font.size = Pt(11)

    for level, size in [(1, 20), (2, 16), (3, 13), (4, 12)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(36)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0F, 0x2A, 0x28)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title_run = title.add_run("E-Barber System")
    title_run.bold = True
    title_run.font.size = Pt(40)
    title_run.font.color.rgb = RGBColor(0x0F, 0x2A, 0x28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run("Część II — Identyfikacja aktorów, przypadki użycia, specyfikacja wymagań")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0x11, 0x5E, 0x59)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(60)
    tag_run = tag.add_run(
        "Dokumentacja inżynierska — transformacja modelu opisowego w model wymagań "
        "na podstawie implementacji aplikacji HairBook Local (Node.js + PostgreSQL)."
    )
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = RGBColor(0x33, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prezentacja: 30.05.2026").bold = True

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document) -> None:
    heading = doc.add_paragraph("Spis treści", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    toc_para = doc.add_paragraph()
    add_field(toc_para, 'TOC \\o "1-3" \\h \\z \\u',
              "Wybierz w Wordzie: prawy klik → Aktualizuj pole, aby wygenerować spis.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_runs(paragraph, text: str) -> None:
    """Render inline markdown (bold/italic/code) into runs."""
    for chunk in INLINE_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(chunk)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header, *body = rows
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for column, text in enumerate(header):
        cell = table.rows[0].cells[column]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:color"), "auto")
        shade.set(qn("w:fill"), "0F766E")
        cell._tc.get_or_add_tcPr().append(shade)

    for row_index, body_row in enumerate(body, start=1):
        for col_index, cell_text in enumerate(body_row):
            if col_index >= len(header):
                break
            cell = table.rows[row_index].cells[col_index]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_runs(para, cell_text.strip())

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    body = "\n".join(lines)
    paragraph = doc.add_paragraph(style="Code Block")
    paragraph.paragraph_format.left_indent = Cm(0.4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F1F5F4")
    paragraph._p.get_or_add_pPr().append(shading)
    paragraph.add_run(body)


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Return (rows, next_index) for a markdown table beginning at start."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = lines[i].strip().strip("|")
        cells = [cell.strip() for cell in row.split("|")]
        rows.append(cells)
        i += 1
    # Drop the alignment row (---|---|...)
    if len(rows) >= 2 and all(set(cell.replace(":", "").replace("-", "").replace(" ", "")) <= {""} for cell in rows[1]):
        rows.pop(1)
    return rows, i


HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")


def render_markdown(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    i = 0
    skip_top_h1 = True
    skip_until_blank_after_top = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            if i < len(lines):
                i += 1
            continue

        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2).strip().strip("*")
            if level == 1 and skip_top_h1:
                # Replaced by the title page
                skip_top_h1 = False
                skip_until_blank_after_top = True
                i += 1
                continue
            level = min(level, 4)
            doc.add_paragraph(content, style=f"Heading {level}")
            i += 1
            continue

        if skip_until_blank_after_top:
            # Skip the title's italic tagline once consumed.
            if stripped.startswith("**") or stripped.startswith("_") or stripped.startswith("*"):
                i += 1
                continue
            if stripped.startswith("Prezentacja"):
                i += 1
                continue
            skip_until_blank_after_top = False

        if stripped.startswith("|"):
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs(paragraph, stripped[2:])
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_runs(paragraph, numbered.group(2))
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_runs(paragraph, stripped)
        i += 1


def main() -> int:
    if not SRC.exists():
        print(f"Source markdown not found: {SRC}", file=sys.stderr)
        return 1

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    ensure_styles(doc)
    add_title_page(doc)
    add_toc(doc)
    render_markdown(doc, SRC.read_text(encoding="utf-8"))

    doc.save(DEST)
    print(f"Saved {DEST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
